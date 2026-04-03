#!/usr/bin/env python3
"""
Geometrie-Arbeitsblatt Generator – Klasse 10
Visualisiert den Zusammenhang zwischen 2D-Flächen und 3D-Körpern:
  Fläche × Höhe → Volumen
"""

import io
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import Polygon, FancyArrowPatch
from mpl_toolkits.mplot3d.art3d import Poly3DCollection
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Farb-Palette
# ---------------------------------------------------------------------------
C2D_FACE   = "#AED6F1"   # Hellblau – 2D-Fläche
C2D_EDGE   = "#1A5276"   # Dunkelblau – 2D-Kante
C3D_TOP    = "#FDEBD0"   # Cremeorange – Deckfläche
C3D_FRONT  = "#F0B27A"   # Orange – Vorderfläche
C3D_SIDE   = "#CA6F1E"   # Dunkelorange – Seitenfläche
C3D_EDGE   = "#784212"   # Braun – 3D-Kante
CH_COLOR   = "#1E8449"   # Grün – Höhe h
CDIM       = "#1A5276"   # Dunkelblau – Beschriftungen
CARROW     = "#7D3C98"   # Violett – Transformationspfeil
CFORMULA   = "#C0392B"   # Rot – Formeln

DPI = 130  # Bildqualität


# ===========================================================================
# Hilfsfunktionen für 3D-Isometrie
# ===========================================================================

def iso(x, y, z, angle_deg=30):
    """Isometrische Projektion: 3D → 2D."""
    a = np.radians(angle_deg)
    xi = x - z * np.cos(a)
    yi = y + z * np.sin(a)
    return xi, yi


def box_faces(dx, dy, dz, ox=0, oy=0):
    """Gibt Polygon-Vertices für die 3 sichtbaren Seiten eines Quaders zurück."""
    pts = {
        "A": iso(ox,      oy,      0),
        "B": iso(ox + dx, oy,      0),
        "C": iso(ox + dx, oy + dy, 0),
        "D": iso(ox,      oy + dy, 0),
        "E": iso(ox,      oy,      dz),
        "F": iso(ox + dx, oy,      dz),
        "G": iso(ox + dx, oy + dy, dz),
        "H": iso(ox,      oy + dy, dz),
    }

    def p(keys):
        return np.array([pts[k] for k in keys])

    front = p(["A", "B", "F", "E"])
    side  = p(["B", "C", "G", "F"])
    top   = p(["E", "F", "G", "H"])
    return front, side, top, pts


def cylinder_patches(r=0.6, h=1.2, cx=0, cy=0, n=60):
    """Gibt die Polygone für Zylinder (Seite + Deckel) zurück (isometrisch)."""
    theta = np.linspace(0, 2 * np.pi, n)

    # Ellipse-Parameter für isometrische Darstellung
    a = r            # Halbachse x
    b = r * 0.35     # Halbachse y (gestaucht)

    # Untere Ellipse
    bx = cx + a * np.cos(theta)
    by = cy + b * np.sin(theta)

    # Obere Ellipse
    tx = bx
    ty = by + h

    # Seite: sichtbare Hälfte (Halbzylinder vorne)
    mask = np.sin(theta) <= 0   # untere Hälfte = Vorderseite
    side_x = np.concatenate([bx[mask], tx[mask][::-1]])
    side_y = np.concatenate([by[mask], ty[mask][::-1]])

    bottom = np.stack([bx, by], axis=1)
    top    = np.stack([tx, ty], axis=1)
    side   = np.stack([side_x, side_y], axis=1)
    return bottom, top, side


def cone_patches(r=0.6, h=1.4, cx=0, cy=0, n=60):
    """Gibt Polygone für einen Kegel zurück."""
    theta = np.linspace(0, 2 * np.pi, n)
    a = r
    b = r * 0.35

    bx = cx + a * np.cos(theta)
    by = cy + b * np.sin(theta)

    tip_x = cx
    tip_y = cy + h

    # Vorderseite
    mask = np.sin(theta) <= 0
    side_x = np.concatenate([bx[mask], [tip_x]])
    side_y = np.concatenate([by[mask], [tip_y]])

    bottom = np.stack([bx, by], axis=1)
    side   = np.stack([side_x, side_y], axis=1)
    return bottom, side


# ===========================================================================
# Einzelne Figurenpaare
# ===========================================================================

def make_figure(draw_fn, figsize=(12, 4.5)):
    """Erstellt eine Matplotlib-Figure, gibt sie als JPEG-Bytes zurück."""
    fig, axes = plt.subplots(1, 3, figsize=figsize,
                             gridspec_kw={"width_ratios": [2, 0.7, 2]})
    fig.patch.set_facecolor("white")
    draw_fn(axes)
    plt.tight_layout(pad=0.5)
    buf = io.BytesIO()
    fig.savefig(buf, format="jpeg", dpi=DPI, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def _arrow_ax(ax):
    """Mittelachse mit Transformationspfeil."""
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.annotate(
        "", xy=(0.85, 0.5), xytext=(0.15, 0.5),
        arrowprops=dict(arrowstyle="-|>", color=CARROW, lw=2.5,
                        mutation_scale=20),
    )
    ax.text(0.5, 0.62, "+ Höhe $h$", ha="center", va="bottom",
            fontsize=11, color=CH_COLOR, fontweight="bold")


def dim_line(ax, x1, y1, x2, y2, label, offset=(0, 0.08), color=CDIM):
    """Zeichnet eine Bemaßungslinie."""
    ax.annotate(
        "", xy=(x2, y2), xytext=(x1, y1),
        arrowprops=dict(arrowstyle="<->", color=color, lw=1.3),
    )
    mx, my = (x1 + x2) / 2 + offset[0], (y1 + y2) / 2 + offset[1]
    ax.text(mx, my, label, ha="center", va="center",
            fontsize=11, color=color, fontweight="bold")


# ------------------------------------------------------------------  1: Quadrat / Würfel
def draw_quadrat_wuerfel(axes):
    ax2, axm, ax3 = axes

    # --- 2D: Quadrat ---
    ax2.set_xlim(-0.3, 1.7)
    ax2.set_ylim(-0.4, 1.6)
    ax2.set_aspect("equal")
    ax2.axis("off")
    sq = mpatches.Rectangle((0.1, 0.1), 1, 1,
                             linewidth=2, edgecolor=C2D_EDGE, facecolor=C2D_FACE)
    ax2.add_patch(sq)
    dim_line(ax2, 0.1, -0.05, 1.1, -0.05, "$a$", offset=(0, -0.12))
    dim_line(ax2, -0.05, 0.1, -0.05, 1.1, "$a$", offset=(-0.18, 0))
    ax2.text(0.6, 1.25, "Quadrat", ha="center", fontsize=13,
             fontweight="bold", color=C2D_EDGE)
    ax2.text(0.6, -0.3, r"$A = a^2$", ha="center", fontsize=14,
             color=CFORMULA, fontweight="bold")

    # --- Pfeil ---
    _arrow_ax(axm)

    # --- 3D: Würfel ---
    ax3.set_xlim(-0.5, 2.0)
    ax3.set_ylim(-0.5, 1.8)
    ax3.set_aspect("equal")
    ax3.axis("off")
    front, side, top, pts = box_faces(1, 1, 1, ox=0.1, oy=0.1)
    ax3.add_patch(Polygon(front, closed=True, fc=C3D_FRONT, ec=C3D_EDGE, lw=1.5))
    ax3.add_patch(Polygon(side,  closed=True, fc=C3D_SIDE,  ec=C3D_EDGE, lw=1.5))
    ax3.add_patch(Polygon(top,   closed=True, fc=C3D_TOP,   ec=C3D_EDGE, lw=1.5))
    # Höhenbeschriftung
    ax3.annotate(
        "", xy=pts["F"], xytext=pts["B"],
        arrowprops=dict(arrowstyle="<->", color=CH_COLOR, lw=1.8),
    )
    mx = (pts["B"][0] + pts["F"][0]) / 2 + 0.12
    my = (pts["B"][1] + pts["F"][1]) / 2
    ax3.text(mx, my, "$a$", fontsize=11, color=CH_COLOR, fontweight="bold",
             ha="left", va="center")
    ax3.text(0.7, 1.55, "Würfel", ha="center", fontsize=13,
             fontweight="bold", color=C3D_EDGE)
    ax3.text(0.7, -0.4, r"$V = a^2 \cdot a = a^3$", ha="center",
             fontsize=14, color=CFORMULA, fontweight="bold")


# ------------------------------------------------------------------  2: Rechteck / Quader
def draw_rechteck_quader(axes):
    ax2, axm, ax3 = axes

    ax2.set_xlim(-0.3, 2.0)
    ax2.set_ylim(-0.5, 1.5)
    ax2.set_aspect("equal")
    ax2.axis("off")
    rect = mpatches.Rectangle((0.1, 0.1), 1.5, 0.9,
                               linewidth=2, edgecolor=C2D_EDGE, facecolor=C2D_FACE)
    ax2.add_patch(rect)
    dim_line(ax2, 0.1, -0.05, 1.6, -0.05, "$a$", offset=(0, -0.12))
    dim_line(ax2, -0.05, 0.1, -0.05, 1.0, "$b$", offset=(-0.18, 0))
    ax2.text(0.85, 1.2, "Rechteck", ha="center", fontsize=13,
             fontweight="bold", color=C2D_EDGE)
    ax2.text(0.85, -0.38, r"$A = a \cdot b$", ha="center", fontsize=14,
             color=CFORMULA, fontweight="bold")

    _arrow_ax(axm)

    ax3.set_xlim(-0.5, 2.5)
    ax3.set_ylim(-0.5, 1.8)
    ax3.set_aspect("equal")
    ax3.axis("off")
    front, side, top, pts = box_faces(1.5, 0.9, 1.0, ox=0.1, oy=0.1)
    ax3.add_patch(Polygon(front, closed=True, fc=C3D_FRONT, ec=C3D_EDGE, lw=1.5))
    ax3.add_patch(Polygon(side,  closed=True, fc=C3D_SIDE,  ec=C3D_EDGE, lw=1.5))
    ax3.add_patch(Polygon(top,   closed=True, fc=C3D_TOP,   ec=C3D_EDGE, lw=1.5))
    ax3.annotate(
        "", xy=pts["F"], xytext=pts["B"],
        arrowprops=dict(arrowstyle="<->", color=CH_COLOR, lw=1.8),
    )
    mx = (pts["B"][0] + pts["F"][0]) / 2 + 0.12
    my = (pts["B"][1] + pts["F"][1]) / 2
    ax3.text(mx, my, "$h$", fontsize=11, color=CH_COLOR, fontweight="bold",
             ha="left", va="center")
    ax3.text(0.95, 1.55, "Quader", ha="center", fontsize=13,
             fontweight="bold", color=C3D_EDGE)
    ax3.text(0.95, -0.4, r"$V = a \cdot b \cdot h$", ha="center",
             fontsize=14, color=CFORMULA, fontweight="bold")


# ------------------------------------------------------------------  3: Kreis / Zylinder
def draw_kreis_zylinder(axes):
    ax2, axm, ax3 = axes

    # --- 2D: Kreis ---
    ax2.set_xlim(-1.4, 1.4)
    ax2.set_ylim(-1.5, 1.5)
    ax2.set_aspect("equal")
    ax2.axis("off")
    circ = plt.Circle((0, 0), 1.0, linewidth=2,
                       edgecolor=C2D_EDGE, facecolor=C2D_FACE)
    ax2.add_patch(circ)
    # Radius-Linie
    ax2.plot([0, 1.0], [0, 0], color=CDIM, lw=1.8)
    ax2.text(0.5, 0.1, "$r$", fontsize=12, color=CDIM, fontweight="bold")
    ax2.text(0, 1.2, "Kreis", ha="center", fontsize=13,
             fontweight="bold", color=C2D_EDGE)
    ax2.text(0, -1.35, r"$A = \pi \cdot r^2$", ha="center", fontsize=14,
             color=CFORMULA, fontweight="bold")

    _arrow_ax(axm)

    # --- 3D: Zylinder ---
    ax3.set_xlim(-1.4, 1.4)
    ax3.set_ylim(-1.0, 2.2)
    ax3.set_aspect("equal")
    ax3.axis("off")
    bottom, top, side = cylinder_patches(r=1.0, h=1.4, cx=0, cy=-0.3)

    ax3.add_patch(Polygon(side,   closed=True, fc=C3D_FRONT, ec=C3D_EDGE, lw=1.5,
                          alpha=0.9))
    ax3.add_patch(Polygon(bottom, closed=True, fc=C3D_SIDE,  ec=C3D_EDGE, lw=1.5,
                          alpha=0.6))
    ax3.add_patch(Polygon(top,    closed=True, fc=C3D_TOP,   ec=C3D_EDGE, lw=1.5))

    # Radius auf Deckel
    ax3.plot([0, 1.0], [1.1, 1.1], color=CDIM, lw=1.8)
    ax3.text(0.5, 1.2, "$r$", fontsize=11, color=CDIM, fontweight="bold")
    # Höhe
    ax3.annotate(
        "", xy=(1.15, 1.1), xytext=(1.15, -0.3),
        arrowprops=dict(arrowstyle="<->", color=CH_COLOR, lw=1.8),
    )
    ax3.text(1.3, 0.4, "$h$", fontsize=11, color=CH_COLOR, fontweight="bold",
             va="center")

    ax3.text(0, 2.05, "Zylinder", ha="center", fontsize=13,
             fontweight="bold", color=C3D_EDGE)
    ax3.text(0, -0.9, r"$V = \pi \cdot r^2 \cdot h$", ha="center",
             fontsize=14, color=CFORMULA, fontweight="bold")


# ------------------------------------------------------------------  4: Dreieck / Prisma
def draw_dreieck_prisma(axes):
    ax2, axm, ax3 = axes

    # --- 2D: Dreieck ---
    ax2.set_xlim(-0.3, 1.8)
    ax2.set_ylim(-0.5, 1.6)
    ax2.set_aspect("equal")
    ax2.axis("off")
    tri_pts = np.array([[0.1, 0.1], [1.5, 0.1], [0.6, 1.2]])
    tri = Polygon(tri_pts, closed=True, linewidth=2,
                  edgecolor=C2D_EDGE, facecolor=C2D_FACE)
    ax2.add_patch(tri)
    # Grundlinie
    dim_line(ax2, 0.1, -0.05, 1.5, -0.05, "$g$", offset=(0, -0.12))
    # Höhe (gestrichelt)
    ax2.plot([0.6, 0.6], [0.1, 1.2], color=CH_COLOR, lw=1.5,
             linestyle="--")
    ax2.text(0.68, 0.6, "$h_A$", fontsize=11, color=CH_COLOR, fontweight="bold")
    ax2.text(0.8, 1.45, "Dreieck", ha="center", fontsize=13,
             fontweight="bold", color=C2D_EDGE)
    ax2.text(0.8, -0.38, r"$A = \dfrac{1}{2} \cdot g \cdot h_A$",
             ha="center", fontsize=13, color=CFORMULA, fontweight="bold")

    _arrow_ax(axm)

    # --- 3D: Dreiecksprisma (isometrisch gezeichnet) ---
    ax3.set_xlim(-0.4, 2.4)
    ax3.set_ylim(-0.5, 1.8)
    ax3.set_aspect("equal")
    ax3.axis("off")

    dz = 0.9   # Tiefe des Prismas

    def ip(x, y, z): return iso(x, y, z, angle_deg=25)

    # Dreieck hinten
    A = np.array(ip(0.1, 0.1, dz))
    B = np.array(ip(1.5, 0.1, dz))
    C = np.array(ip(0.6, 1.2, dz))
    # Dreieck vorne
    D = np.array(ip(0.1, 0.1, 0))
    E = np.array(ip(1.5, 0.1, 0))
    F = np.array(ip(0.6, 1.2, 0))

    # Untere Rechteckfläche (D-E-B-A)
    bottom_face = np.array([D, E, B, A])
    # Vordere Dreiecksfläche
    front_face = np.array([D, E, F])
    # Seitliche Fläche links (D-F-C-A)
    left_face = np.array([D, F, C, A])

    ax3.add_patch(Polygon(bottom_face, closed=True, fc=C3D_FRONT,
                          ec=C3D_EDGE, lw=1.5))
    ax3.add_patch(Polygon(left_face,  closed=True, fc=C3D_SIDE,
                          ec=C3D_EDGE, lw=1.5))
    ax3.add_patch(Polygon(front_face, closed=True, fc=C3D_TOP,
                          ec=C3D_EDGE, lw=1.5))
    # Hintere Kanten sichtbar machen
    for start, end in [(A, B), (A, C), (B, C)]:
        ax3.plot([start[0], end[0]], [start[1], end[1]],
                 color=C3D_EDGE, lw=1.2, linestyle="--", alpha=0.6)

    # Tiefe/Höhe beschriften
    mid_DE = (D + E) / 2
    mid_AB = (A + B) / 2
    ax3.annotate(
        "", xy=tuple(mid_AB), xytext=tuple(mid_DE),
        arrowprops=dict(arrowstyle="<->", color=CH_COLOR, lw=1.8),
    )
    mid = (mid_DE + mid_AB) / 2
    ax3.text(mid[0], mid[1] - 0.15, "$h$", fontsize=11,
             color=CH_COLOR, fontweight="bold", ha="center")

    ax3.text(0.9, 1.6, "Dreiecksprisma", ha="center", fontsize=13,
             fontweight="bold", color=C3D_EDGE)
    ax3.text(0.9, -0.4, r"$V = \dfrac{1}{2} \cdot g \cdot h_A \cdot h$",
             ha="center", fontsize=13, color=CFORMULA, fontweight="bold")


# ------------------------------------------------------------------  5: Kreis / Kegel
def draw_kreis_kegel(axes):
    ax2, axm, ax3 = axes

    # --- 2D: Kreis ---
    ax2.set_xlim(-1.4, 1.4)
    ax2.set_ylim(-1.5, 1.5)
    ax2.set_aspect("equal")
    ax2.axis("off")
    circ = plt.Circle((0, 0), 1.0, linewidth=2,
                       edgecolor=C2D_EDGE, facecolor=C2D_FACE)
    ax2.add_patch(circ)
    ax2.plot([0, 1.0], [0, 0], color=CDIM, lw=1.8)
    ax2.text(0.5, 0.1, "$r$", fontsize=12, color=CDIM, fontweight="bold")
    ax2.text(0, 1.2, "Kreis", ha="center", fontsize=13,
             fontweight="bold", color=C2D_EDGE)
    ax2.text(0, -1.35, r"$A = \pi \cdot r^2$", ha="center", fontsize=14,
             color=CFORMULA, fontweight="bold")

    _arrow_ax(axm)

    # --- 3D: Kegel ---
    ax3.set_xlim(-1.4, 1.4)
    ax3.set_ylim(-1.0, 2.2)
    ax3.set_aspect("equal")
    ax3.axis("off")
    bottom, side = cone_patches(r=1.0, h=1.6, cx=0, cy=-0.3)

    ax3.add_patch(Polygon(side,   closed=True, fc=C3D_FRONT, ec=C3D_EDGE,
                          lw=1.5, alpha=0.9))
    ax3.add_patch(Polygon(bottom, closed=True, fc=C3D_SIDE,  ec=C3D_EDGE,
                          lw=1.5, alpha=0.6))

    # Radius
    ax3.plot([0, 1.0], [-0.3, -0.3], color=CDIM, lw=1.8)
    ax3.text(0.5, -0.22, "$r$", fontsize=11, color=CDIM, fontweight="bold")
    # Höhe
    ax3.annotate(
        "", xy=(1.2, 1.3), xytext=(1.2, -0.3),
        arrowprops=dict(arrowstyle="<->", color=CH_COLOR, lw=1.8),
    )
    ax3.text(1.35, 0.5, "$h$", fontsize=11, color=CH_COLOR, fontweight="bold",
             va="center")

    ax3.text(0, 2.05, "Kegel", ha="center", fontsize=13,
             fontweight="bold", color=C3D_EDGE)
    ax3.text(0, -0.9, r"$V = \dfrac{1}{3} \cdot \pi \cdot r^2 \cdot h$",
             ha="center", fontsize=14, color=CFORMULA, fontweight="bold")


# ===========================================================================
# Word-Dokument
# ===========================================================================

SHAPES = [
    ("Quadrat → Würfel",
     r"Das Quadrat hat die Fläche A = a². Strecken wir es in die dritte "
     r"Dimension um die Höhe a, entsteht der Würfel mit V = a³.",
     draw_quadrat_wuerfel),

    ("Rechteck → Quader",
     r"Das Rechteck mit A = a · b wird durch Hinzufügen der Höhe h zum "
     r"Quader: V = a · b · h.",
     draw_rechteck_quader),

    ("Kreis → Zylinder",
     r"Der Kreis mit A = π · r² wird durch eine Höhe h zum Zylinder: "
     r"V = π · r² · h.",
     draw_kreis_zylinder),

    ("Dreieck → Dreiecksprisma",
     r"Das Dreieck mit A = ½ · g · h_A wird durch Hinzufügen der Tiefe h "
     r"zum Dreiecksprisma: V = ½ · g · h_A · h.",
     draw_dreieck_prisma),

    ("Kreis → Kegel",
     r"Beim Kegel kommt zum Kreisboden noch der Faktor ⅓ hinzu, weil der "
     r"Körper sich zur Spitze verjüngt: V = ⅓ · π · r² · h.",
     draw_kreis_kegel),
]


def set_heading(paragraph, level=1):
    """Setzt Überschrift-Stil."""
    paragraph.style = f"Heading {level}"


def add_rule(doc):
    """Horizontale Linie als Trenner."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def build_document(output_path="geometrie_arbeitsblatt.docx"):
    doc = Document()

    # ---- Seite einrichten (A4 Querformat) ----
    section = doc.sections[0]
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    # ---- Titel ----
    title = doc.add_heading("Geometrie: Vom Flächeninhalt zum Volumen", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        "Klasse 10  ·  Wie entsteht aus einer 2D-Figur ein 3D-Körper?"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    # ---- Einleitung ----
    intro = doc.add_paragraph(
        "Prinzip: Jeder räumliche Körper entsteht, indem man eine ebene "
        "Grundfigur um eine Hoehe h in den Raum 'schiebt' (Prisma, Zylinder) "
        "oder zuspitzt (Pyramide, Kegel). "
        "Deshalb gilt immer: V = Grundfläche × Höhe  (bei Prismen und Zylindern) "
        "bzw. V = ⅓ × Grundfläche × Höhe  (bei Pyramiden und Kegeln)."
    )
    intro.style.font.size = Pt(11)

    # ---- Figurenpaare ----
    for title_text, explanation, draw_fn in SHAPES:
        add_rule(doc)
        h2 = doc.add_heading(title_text, level=2)
        h2.runs[0].font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

        p_expl = doc.add_paragraph(explanation)
        p_expl.runs[0].font.size = Pt(11)

        # Figur zeichnen & einfügen
        buf = make_figure(draw_fn)
        pic_p = doc.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic_p.add_run()
        run.add_picture(buf, width=Inches(9.5))

    # ---- Zusammenfassung ----
    add_rule(doc)
    doc.add_heading("Zusammenfassung der Formeln", level=2)

    summary_data = [
        ("Figur/Körper", "Fläche A", "Volumen V"),
        ("Quadrat / Würfel",           "a²",              "a³"),
        ("Rechteck / Quader",          "a · b",           "a · b · h"),
        ("Kreis / Zylinder",           "π · r²",          "π · r² · h"),
        ("Dreieck / Dreiecksprisma",   "½ · g · h_A",     "½ · g · h_A · h"),
        ("Kreis / Kegel",              "π · r²",          "⅓ · π · r² · h"),
    ]

    table = doc.add_table(rows=len(summary_data), cols=3)
    table.style = "Light Shading Accent 1"
    for r_idx, row_data in enumerate(summary_data):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = cell_text
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(11)
            if r_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.save(output_path)
    print(f"Dokument gespeichert: {output_path}")
    return output_path


if __name__ == "__main__":
    build_document()
